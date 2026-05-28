"""
Gera arquivo DOCX com mensagens de abordagem personalizadas por lead
Recebe JSON com dados + gargalos de cada lead
"""

import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


BERRY = RGBColor(0x6B, 0x21, 0x4E)
ROSE  = RGBColor(0xC0, 0x6B, 0x8A)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = BERRY if level == 1 else ROSE
    return p


def add_label(doc, label, value, italic_value=False):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.color.rgb = DARK
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.italic = italic_value
    r2.font.size = Pt(11)
    r2.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(2)
    return p


def add_message_box(doc, titulo, mensagem):
    doc.add_paragraph()
    p_titulo = doc.add_paragraph()
    r = p_titulo.add_run(f"  {titulo}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BERRY
    p_titulo.paragraph_format.space_after = Pt(0)

    p_msg = doc.add_paragraph()
    p_msg.paragraph_format.left_indent = Inches(0.2)
    p_msg.paragraph_format.space_before = Pt(2)
    r_msg = p_msg.add_run(mensagem)
    r_msg.font.size = Pt(11)
    r_msg.font.color.rgb = DARK


def gerar_docx(leads: list[dict], output_path: str):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Scripts de Abordagem Personalizados")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BERRY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} · Clube Divos da IA")
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY

    doc.add_paragraph()

    for i, lead in enumerate(leads, 1):
        # Separador entre leads
        if i > 1:
            doc.add_paragraph("─" * 60).paragraph_format.space_before = Pt(12)

        add_heading(doc, f"{i}. {lead['nome']}", level=1)

        if lead.get('telefone'):
            add_label(doc, "📞 Telefone:", lead['telefone'])
        if lead.get('endereco'):
            add_label(doc, "📍 Endereço:", lead['endereco'])
        if lead.get('avaliacao'):
            reviews = f" ({lead.get('reviews', '')} avaliações)" if lead.get('reviews') else ''
            add_label(doc, "⭐ Avaliação:", f"{lead['avaliacao']}{reviews}")
        if lead.get('site'):
            add_label(doc, "🌐 Site:", lead['site'])

        # Gargalos encontrados
        if lead.get('gargalos'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            r = p.add_run("🔍 Gargalos identificados:")
            r.bold = True
            r.font.color.rgb = ROSE
            r.font.size = Pt(11)

            for g in lead['gargalos']:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.paragraph_format.left_indent = Inches(0.3)
                rr = bullet.add_run(g)
                rr.font.size = Pt(11)
                rr.font.color.rgb = GRAY

        doc.add_paragraph()

        # Mensagem principal
        if lead.get('mensagem_whatsapp'):
            add_heading(doc, "💬 Mensagem WhatsApp / Instagram DM", level=2)
            add_message_box(doc, "Primeiro contato:", lead['mensagem_whatsapp'])

        if lead.get('mensagem_followup'):
            add_message_box(doc, "Follow-up (2-3 dias depois):", lead['mensagem_followup'])

        if lead.get('email_assunto'):
            doc.add_paragraph()
            add_heading(doc, "📧 E-mail", level=2)
            add_label(doc, "Assunto:", lead['email_assunto'])
            if lead.get('email_corpo'):
                add_message_box(doc, "Corpo:", lead['email_corpo'])

        doc.add_paragraph()

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python3 gerar_mensagens.py leads_com_mensagens.json")
        sys.exit(1)

    with open(sys.argv[1]) as f:
        leads = json.load(f)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output = str(Path(sys.argv[1]).parent / f"mensagens_{timestamp}.docx")
    gerar_docx(leads, output)
    print(output)
