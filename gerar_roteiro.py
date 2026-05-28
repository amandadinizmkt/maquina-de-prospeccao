"""
Gera o roteiro da aula "Máquina de Prospecção com IA" em DOCX
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BERRY   = RGBColor(0x74, 0x1E, 0x31)
ROSE    = RGBColor(0xC4, 0x90, 0x8E)
DARK    = RGBColor(0x1A, 0x1A, 0x1A)
GRAY    = RGBColor(0x55, 0x55, 0x55)
GREEN   = RGBColor(0x16, 0x61, 0x34)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT   = RGBColor(0xF5, 0xF5, 0xF5)


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = BERRY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = ROSE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def add_body(doc, text, color=None, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARK
    run.italic = italic
    run.bold = bold
    p.paragraph_format.space_after = Pt(3)
    return p


def add_step(doc, icon, label, desc=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{icon}  ")
    r1.font.size = Pt(11)
    r2 = p.add_run(label)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    if desc:
        r3 = p.add_run(f"  — {desc}")
        r3.font.size = Pt(10)
        r3.font.color.rgb = GRAY


def add_cmd(doc, text):
    """Bloco de comando estilo terminal."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {text}  ")
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xA8, 0xFF, 0xC0)
    run.font.bold = True
    # Fundo escuro via shading no parágrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0D1117')
    pPr.append(shd)


def add_nota(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)
    run = p.add_run(f"⚠️  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x85, 0x60, 0x00)


def add_divider(doc):
    p = doc.add_paragraph("─" * 70)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)


def add_summary_table(doc, rows):
    table = doc.add_table(rows=len(rows)+1, cols=4)
    table.style = 'Table Grid'

    headers = ['Parte', 'Conteúdo', 'Tempo', 'Tipo']
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, '741E31')

    colors = {'Setup': 'D6EAF8', 'Demo': 'D5F5E3', 'Deploy': 'FAD7A0', 'Encerramento': 'F9EBEA'}
    for i, row_data in enumerate(rows):
        row = table.rows[i+1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
        tipo = row_data[3]
        bg = colors.get(tipo, 'FFFFFF')
        for cell in row.cells:
            set_cell_bg(cell, bg)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── CAPA ──────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Máquina de Prospecção com IA")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BERRY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Roteiro da Aula · Clube Divos da IA")
    r2.font.size = Pt(12); r2.font.color.rgb = ROSE

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2.add_run("Duração total: ~75 minutos")
    r3.font.size = Pt(10); r3.font.color.rgb = GRAY
    r3.italic = True

    doc.add_paragraph()

    # ── TABELA RESUMO ─────────────────────────────────────────────────────────
    add_heading(doc, "Visão geral da aula", level=2)
    add_summary_table(doc, [
        ('Parte 1', 'Setup ao vivo — clonar + instalar', '20 min', 'Setup'),
        ('Parte 2', 'Demo ao vivo — buscar, escolher, pesquisar, mensagens', '45 min', 'Demo'),
        ('Parte 3', 'Encerramento + tarefa da semana', '10 min', 'Encerramento'),
    ])

    doc.add_paragraph()
    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PARTE 1 — SETUP
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PARTE 1 — Setup ao vivo  [~20 min]")

    # Etapa 1
    add_heading(doc, "Etapa 1 — Abertura e contexto  [3 min]", level=2)
    add_step(doc, "🖥", "Abrir os slides", "Slide 01 — capa")
    add_step(doc, "✅", "Mostrar o Slide 02", "o problema da prospecção manual")
    add_step(doc, "✅", "Mostrar o Slide 03", "o fluxo que vamos montar hoje")
    add_step(doc, "✅", "Mostrar o Slide 04", "as ferramentas que vamos usar")

    # Etapa 2
    add_heading(doc, "Etapa 2 — Clonar o repositório  [5 min]", level=2)
    add_step(doc, "🖥", "Abrir o Claude Code ao vivo")
    add_step(doc, "✅", "Ditar o comando para as alunas")
    add_cmd(doc, '"Clona o repositório github.com/amandadinizmkt/maquina-de-prospeccao e roda o instalar.py"')
    add_step(doc, "🖥", "Mostrar o clone acontecendo no terminal")
    add_nota(doc, "Se alguma aluna não tiver o Claude Code, orienta a instalar em claude.ai/code — é o único pré-requisito.")

    # Etapa 3
    add_heading(doc, "Etapa 3 — Instalação  [7 min]", level=2)
    add_step(doc, "🖥", "Mostrar o instalar.py rodando")
    add_body(doc, "Saída esperada no terminal:")
    add_cmd(doc, "✅ Python 3.14 — ok")
    add_cmd(doc, "✅ Scrapling 0.4.8 — instalado")
    add_cmd(doc, "✅ Playwright Python — instalado")
    add_cmd(doc, "✅ python-docx — instalado")
    add_cmd(doc, "✅ Tudo instalado! Pronto para prospectar.")
    add_step(doc, "✅", "Mostrar o Slide 05", "setup completo")
    add_nota(doc, "O Playwright pode demorar mais na primeira vez — fica uns 2-3 minutos baixando os navegadores. Normal.")

    # Etapa 4
    add_heading(doc, "Etapa 4 — Explicar as ferramentas  [5 min]", level=2)
    add_step(doc, "✅", "Mostrar o Slide 04", "explicar o que cada ferramenta faz")
    add_body(doc, "Pontos a cobrir:")
    add_step(doc, "•", "Scrapling", "entra no Google Maps como se fosse um humano")
    add_step(doc, "•", "Playwright", "abre o Instagram e sites reais para pesquisar os leads")
    add_step(doc, "•", "python-docx", "gera o arquivo Word com as mensagens")
    add_step(doc, "•", "Tudo gratuito", "zero custo além do Claude Code")

    doc.add_paragraph()
    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PARTE 2 — DEMO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PARTE 2 — Demo ao vivo  [~45 min]")

    # Etapa 5
    add_heading(doc, "Etapa 5 — Buscando os leads  [10 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 06", "buscar leads")
    add_step(doc, "🖥", "No Claude Code, digitar ao vivo:")
    add_cmd(doc, '"Busca clínicas de estética em São Paulo"')
    add_body(doc, "O que mostrar enquanto roda (~1 minuto de espera):")
    add_step(doc, "•", "O Google Maps abrindo automaticamente em segundo plano")
    add_step(doc, "•", "Os resultados aparecendo no terminal")
    add_step(doc, "•", "A planilha CSV abrindo no Excel/Numbers")
    add_nota(doc, "Escolha o nicho da demo com antecedência e teste antes da aula. Cidades grandes retornam mais resultados.")

    # Etapa 6
    add_heading(doc, "Etapa 6 — Apresentando a planilha  [5 min]", level=2)
    add_step(doc, "🖥", "Mostrar a planilha aberta com as colunas")
    add_step(doc, "✅", "Mostrar Slide 07", "como ler os dados")
    add_body(doc, "Colunas a destacar:")
    add_step(doc, "•", "Avaliação + Reviews", "combinado indica reputação e volume")
    add_step(doc, "•", "Telefone", "pronto para abordar via WhatsApp")
    add_step(doc, "•", "Endereço", "contexto geográfico para personalizar a mensagem")
    add_nota(doc, "Peça para as alunas apontarem quais levariam para pesquisa. Cria engajamento antes de mostrar a próxima etapa.")

    # Etapa 7
    add_heading(doc, "Etapa 7 — Qualificação com Claude (bônus)  [5 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 11", "prompt de qualificação")
    add_step(doc, "🖥", "Usar o prompt de qualificação ao vivo:")
    add_cmd(doc, '"Analisa esses leads e me diz quais são os 5 com mais potencial..."')
    add_body(doc, "Mostrar o Claude devolvendo o ranking com score e justificativa.")

    # Etapa 8
    add_heading(doc, "Etapa 8 — Escolhendo os leads para pesquisar  [3 min]", level=2)
    add_step(doc, "🖥", "Ditar o comando ao vivo:")
    add_cmd(doc, '"Quero pesquisar os leads 3, 7, 12, 15 e 20"')
    add_body(doc, "ou pedir para o Claude escolher:")
    add_cmd(doc, '"Você escolhe os 5 com mais potencial para gestão de redes"')

    # Etapa 9
    add_heading(doc, "Etapa 9 — Pesquisa de gargalos  [10 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 08", "pesquisa de gargalos")
    add_step(doc, "🖥", "Mostrar o Claude visitando o Instagram de cada lead ao vivo")
    add_body(doc, "O que mostrar:")
    add_step(doc, "•", "Claude abrindo o Google e o Instagram")
    add_step(doc, "•", "Dados reais: seguidores, posts, bio")
    add_step(doc, "•", "Gargalos identificados com linguagem direta")
    add_body(doc, "Exemplos reais para destacar:")
    add_step(doc, "→", "\"2.006 posts com 11K seguidores — baixa taxa de conversão\"")
    add_step(doc, "→", "\"Seguindo 2.975 contas — prejudica o alcance orgânico\"")
    add_step(doc, "→", "\"Nota 2,5 no Maps — crise de reputação visível para todos\"")
    add_nota(doc, "Esse é o momento mais impactante da aula. As alunas veem o Claude encontrando informações reais que elas usariam numa abordagem. Deixa o silêncio render.")

    # Etapa 10
    add_heading(doc, "Etapa 10 — Gerando as mensagens  [10 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 09", "mensagens personalizadas")
    add_step(doc, "🖥", "Ditar o comando ao vivo:")
    add_cmd(doc, '"Gera as mensagens personalizadas. Meu nome é [nome], trabalho com gestão de redes para [nicho]."')
    add_step(doc, "🖥", "Abrir o DOCX gerado e mostrar ao vivo")
    add_body(doc, "Destacar nas mensagens:")
    add_step(doc, "•", "Cada mensagem menciona algo específico do negócio")
    add_step(doc, "•", "O gargalo real aparece naturalmente na abordagem")
    add_step(doc, "•", "Tom humano, sem parecer automático")
    add_step(doc, "•", "Termina com pergunta aberta — não empurra a venda")
    add_nota(doc, "Compare com a mensagem genérica \"Oi, faço gestão de redes, vamos conversar?\" — o contraste é o ponto mais forte da aula.")

    # Etapa 11
    add_heading(doc, "Etapa 11 — Fluxo completo  [2 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 10", "fluxo completo com tempos")
    add_body(doc, "Enfatizar: do zero à mensagem personalizada em ~15 minutos.")

    doc.add_paragraph()
    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PARTE 3 — ENCERRAMENTO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PARTE 3 — Encerramento  [~10 min]")

    # Etapa 12
    add_heading(doc, "Etapa 12 — Tarefa da semana  [5 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 12", "tarefa")
    add_body(doc, "Tarefa:")
    add_step(doc, "1.", "Instalar a máquina de prospecção")
    add_step(doc, "2.", "Buscar 30 leads do nicho principal")
    add_step(doc, "3.", "Pesquisar os 5 melhores e enviar pelo menos 3 mensagens reais")
    add_body(doc, "Lembrar de trazer os resultados para a próxima aula.", color=ROSE, italic=True)

    # Etapa 13
    add_heading(doc, "Etapa 13 — Encerramento  [5 min]", level=2)
    add_step(doc, "✅", "Mostrar Slide 13", "encerramento")
    add_body(doc, "Preview da próxima aula:")
    add_step(doc, "→", "Como Fechar Mais com IA — Da Proposta ao Sim")
    add_step(doc, "→", "Sistema de geração automática de proposta personalizada")
    add_step(doc, "→", "Scripts de objeção e decisão de próximo passo")

    doc.add_paragraph()
    add_divider(doc)

    # ── PROMPTS DE REFERÊNCIA ─────────────────────────────────────────────────
    add_heading(doc, "Prompts de referência  (para usar durante a aula)", level=2)

    prompts = [
        ("Busca de leads", '"Busca [nicho] em [cidade]"'),
        ("Qualificação", '"Analisa esses leads e me diz os 5 com mais potencial para contratar gestão de redes"'),
        ("Escolha", '"Quero pesquisar os leads 3, 7 e 12"'),
        ("Mensagens", '"Gera as mensagens. Meu nome é [nome], trabalho com [serviço] para [nicho]"'),
        ("Refinamento", '"Essa mensagem ficou genérica. Reescreve mencionando o gargalo de [gargalo]"'),
        ("Abrir planilha", '"Abre a planilha no Finder"'),
    ]

    for label, cmd in prompts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = ROSE
        r2 = p.add_run(cmd)
        r2.font.name = 'Courier New'; r2.font.size = Pt(10); r2.font.color.rgb = DARK

    doc.save('/Users/amandaroddiniz/aula-prospeccao/roteiro-maquina-de-prospeccao.docx')
    print("✅ Roteiro salvo!")


if __name__ == "__main__":
    main()
